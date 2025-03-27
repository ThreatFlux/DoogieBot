import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import json
import markdown
import pypdf  # Replaced deprecated PyPDF2
from docx import Document as DocxDocument
import yaml

from app.models.document import DocumentType

class DocumentParser:
    """
    Parser for different document types.
    """
    
    @staticmethod
    def parse_document(file_path: str, doc_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a document based on its type and return the text content and metadata.
        """
        if not os.path.exists(file_path) and doc_type != DocumentType.MANUAL:
            raise FileNotFoundError(f"File not found: {file_path}")
        
        content = ""
        metadata = {}
        
        if doc_type == DocumentType.PDF:
            content, metadata = DocumentParser.parse_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            content, metadata = DocumentParser.parse_docx(file_path)
        elif doc_type == DocumentType.MARKDOWN:
            content, metadata = DocumentParser.parse_markdown(file_path)
        elif doc_type == DocumentType.RST:
            content, metadata = DocumentParser.parse_rst(file_path)
        elif doc_type == DocumentType.TEXT:
            content, metadata = DocumentParser.parse_text(file_path)
        elif doc_type == DocumentType.JSON:
            content, metadata = DocumentParser.parse_json(file_path)
        elif doc_type == DocumentType.JSONL:
            content, metadata = DocumentParser.parse_jsonl(file_path)
        elif doc_type == DocumentType.YAML or doc_type == DocumentType.YML:
            content, metadata = DocumentParser.parse_yaml(file_path)
        elif doc_type == DocumentType.MANUAL:
            # For manual documents, the file_path is actually the content
            content = file_path
            metadata = {"type": "manual"}
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
        
        return content, metadata
    
    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a PDF file and return the text content and metadata.
        """
        text = ""
        metadata = {
            "pages": 0,
            "title": None,
            "author": None,
            "subject": None,
            "keywords": None
        }
        
        with open(file_path, "rb") as f:
            pdf = pypdf.PdfReader(f)
            metadata["pages"] = len(pdf.pages)
            
            # Extract document info
            if pdf.metadata:
                metadata["title"] = pdf.metadata.get('/Title')
                metadata["author"] = pdf.metadata.get('/Author')
                metadata["subject"] = pdf.metadata.get('/Subject')
                metadata["keywords"] = pdf.metadata.get('/Keywords')
            
            # Extract text from each page
            for page in pdf.pages:
                text += page.extract_text() + "\n\n"
        
        return text, metadata
    
    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a DOCX file and return the text content and metadata.
        """
        doc = DocxDocument(file_path)
        
        # Extract text
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract metadata
        metadata = {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified,
            "paragraphs": len(doc.paragraphs)
        }
        
        return text, metadata
    
    @staticmethod
    def parse_markdown(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a Markdown file and return the text content and metadata.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        # Convert to HTML (for potential future use)
        html = markdown.markdown(text)
        
        metadata = {
            "format": "markdown",
            "html_length": len(html)
        }
        
        return text, metadata
    
    @staticmethod
    def parse_rst(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse an RST file and return the text content and metadata.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        metadata = {
            "format": "rst"
        }
        
        return text, metadata
    
    @staticmethod
    def parse_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a text file and return the text content and metadata.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        metadata = {
            "format": "text",
            "lines": text.count("\n") + 1
        }
        
        return text, metadata
    
    @staticmethod
    def parse_json(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a JSON file and return the text content and metadata,
        including the parsed JSON structure in the metadata.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Convert JSON to a compact string for the main 'content' field
        # Using separators=(',', ':') creates the most compact valid JSON string
        text = json.dumps(data, separators=(',', ':'))
        
        metadata = {
            "format": "json",
            "keys": list(data.keys()) if isinstance(data, dict) else [],
            "parsed_json": data  # Add the raw parsed data here
        }
        
        return text, metadata
    
    @staticmethod
    def parse_jsonl(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a JSONL file and return the text content and metadata.
        """
        lines = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    data = json.loads(line.strip())
                    lines.append(data)
                except json.JSONDecodeError:
                    continue
        
        # Convert to string for indexing
        text = "\n".join([json.dumps(line) for line in lines])
        
        metadata = {
            "format": "jsonl",
            "lines": len(lines)
        }
        
        return text, metadata
        
    @staticmethod
    def parse_yaml(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse a YAML file and return the text content and metadata.
        """
        with open(file_path, "r", encoding="utf-8") as f:
            try:
                # Load YAML content
                data = yaml.safe_load(f)
                
                # Convert YAML to string for indexing
                if isinstance(data, dict):
                    text = yaml.dump(data, default_flow_style=False)
                    keys = list(data.keys())
                elif isinstance(data, list):
                    text = yaml.dump(data, default_flow_style=False)
                    keys = []
                    for item in data:
                        if isinstance(item, dict):
                            keys.extend(list(item.keys()))
                    # Remove duplicates
                    keys = list(set(keys))
                else:
                    text = str(data)
                    keys = []
                
                metadata = {
                    "format": "yaml",
                    "keys": keys
                }
                
                return text, metadata
            except yaml.YAMLError as e:
                # If there's an error parsing the YAML, return the raw content
                f.seek(0)
                text = f.read()
                metadata = {
                    "format": "yaml",
                    "parse_error": str(e)
                }
                return text, metadata